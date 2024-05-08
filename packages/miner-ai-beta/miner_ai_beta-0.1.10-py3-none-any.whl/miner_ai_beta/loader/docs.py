import os
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from tqdm import tqdm


def IndexFromDocs(folder_path: str, embeddings, vectorstore, chunk_size: int = 2000):
    """
    ### Load documents from folder
    This function loads documents from a folder and creates a Fass index (db).\n
    - folder_path: Path to folder containing PDFs\n
    - chunk_size: Chunk size in characters (default 2000)\n
    - embeddings: Embeddings object from langchain.embeddings\n
    - vectorstore: could be FAISS or ChromaDB\n
    """

    # Assuming your PDFs are in a folder named 'test_pdfs'
    doc_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]

    documents = []
    # Process each PDF file and split into chunks of 24000 characters
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=0, length_function=len, is_separator_regex=False) # FIXED

    super_text = ""
    documents = []
    # Process each PDF file
    for doc_file in tqdm(doc_files, "Parsing docs"):
        doc = docx.Document(os.path.join(folder_path, doc_file))
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        super_text += '\n'.join(fullText)   
        print(f"Doc '{doc_file}' ingested")

        texts = text_splitter.split_text(super_text)

        for text in texts:
            document = Document(page_content=text, metadata={"title": doc_file, "type": "docx", "tags": ["document"]})
            documents.append(document)

    # Store Embeddings in FAISS
    for i in tqdm(range(0,1), "Creating db from documents"):
        db = vectorstore.from_documents(documents, embeddings)
    return db
