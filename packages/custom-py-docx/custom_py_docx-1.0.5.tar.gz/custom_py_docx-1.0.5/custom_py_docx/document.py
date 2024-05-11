import os
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer
import plotly.express as px
from typing import List, Optional, Union
import random
import colorsys
import plotly.graph_objs as go

class CreateDocx:
    def __init__(self, doc_path='', doc_name='new_doc', font_name='Tahoma', font_size=11) -> None:
        self.doc = Document()
        self.doc_name = doc_name
        self.doc_path = doc_path
        self.font_name = font_name
        self.font_size = font_size
        self.set_font()
    
    def set_font(self):
        font = self.doc.styles['Normal'].font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        
    def save_doc(self):
        self.doc.save(f'{self.doc_path}{self.doc_name}.docx')
    
    def delete_doc(doc_path):
        os.remove(doc_path)
    
    def concat_docs(self, doc1_path, doc2_path, new_doc_path, new_doc_name):
        doc1 = Document(doc1_path)
        doc2 = Document(doc2_path)
        composer = Composer(doc1)
        composer.append(doc2)
        composer.save(f'{new_doc_path}/{new_doc_name}')
    
    def break_page(self):
        self.doc.add_page_break()
   
    def add_heading(self, text, level):
        self.doc.add_heading(text, level)
    
    def add_image(self, image_path_name, widht_inches=3.0, height_inches=3.0):
        self.doc.add_picture(image_path_name, width=Inches(widht_inches), height=Inches(height_inches))
        last_paragraph = self.doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_paragraph(self, text: str, bold=False, italic=False, set_color={"r": 0, "g": 0, "b": 0}):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(set_color['r'], set_color['g'], set_color['b'])
        return paragraph        

    def concat_on_paragraph(self, paragraph, text, bold=False, italic=False, set_color={"r": 0, "g": 0, "b": 0}):
        run = paragraph.add_run(text)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(set_color['r'], set_color['g'], set_color['b'])

    def create_table(self, table_title, table_title_level, style='Medium Shading 1', column_names=[], rows=()):
        self.doc.add_heading(table_title, table_title_level)
        table = self.doc.add_table(rows=1, cols=len(column_names), style=style)
        row = table.rows[0].cells
       
        for i in range(len(column_names)):
            row[i].text = column_names[i]
        
        row_len = len(rows[0])
        for row_data in rows:
            row = table.add_row().cells
            row_index = 0
            while row_index < row_len:
                row[row_index].text = str(row_data[row_index])
                row_index += 1

    def add_pizza_plot(self, x_values, y_values, width=450, height=400, plot_name=''):
        fig = px.pie(names=x_values, values=y_values, width=width, height=height)
        fig.update_traces(textposition='inside', textinfo='percent+label')
        
    
        img_bytes = fig.to_image(format="png", engine="kaleido")
        with open(f"{plot_name}.png", "wb") as file:
            file.write(img_bytes)
        self.doc.add_picture(f"{plot_name}.png")
        # os.remove(f"{plot_name}.png")

    def add_bar_plot(
        self,
        x_values: List[Union[str, int, float]], 
        y_values: List[Union[int, float]], 
        colors: Optional[List[str]] = None, 
        y_title: Optional[str] = None, 
        x_title: Optional[str] = None, 
        plot_name: Optional[str] = "plot", 
        random_colors: Optional[bool] = False, 
        base_color: Optional[str] = None
    ) -> None:
        
        fig = go.Figure()
        
        # set the marker color for the bar graph
        if colors is not None:
            fig.add_trace(
                go.Bar(x=x_values, y=y_values, marker=dict(color=colors))
            )
        else:
            if random_colors:
                # generate random colors based on base_color if provided
                num_bars = len(x_values)
                if base_color is not None:
                    # convert the base color to HSL
                    base_hsl = colorsys.rgb_to_hls(*tuple(int(base_color[i:i+2], 16) for i in (0, 2, 4)))
                    # vary the lightness of the base color
                    colors = [f"rgb{tuple(map(int, colorsys.hls_to_rgb(base_hsl[0], max(0, min(1, base_hsl[1] + random.uniform(-0.2, 0.2))), max(0, min(1, base_hsl[2] + random.uniform(-0.2, 0.2))))))}"
                            for _ in range(num_bars)]
                else:
                    colors = [f"rgb({random.randint(0, 255)}, {random.randint(0, 255)}, {random.randint(0, 255)})"
                            for _ in range(num_bars)]
                fig.add_trace(
                    go.Bar(x=x_values, y=y_values, marker=dict(color=colors))
                )
            else:
                fig.add_trace(
                    go.Bar(x=x_values, y=y_values)
                )
                

        # Set the figure layout
        fig.update_layout(
            xaxis_title=x_title,
            yaxis_title=y_title,
            width=500, # Set the width of the figure
            height=500, # Set the height of the figure
            margin=dict(l=30, r=30, t=30, b=30), # Set the margin
            font=dict(size=12) # Set the font size of the text in the figure
        )
        
        # save the plot as a png image
        img_bytes = fig.to_image(format="png", engine="kaleido")
        with open(f"{plot_name}.png", "wb") as file:
            file.write(img_bytes)
        self.doc.add_picture(f"{plot_name}.png")