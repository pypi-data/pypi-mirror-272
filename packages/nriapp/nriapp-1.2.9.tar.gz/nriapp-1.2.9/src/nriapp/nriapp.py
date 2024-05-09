#!/usr/bin/python
#this is a test
import os, sys
import time
import requests, json, pickle, ast
import queue
import re

sys.path.append("..")

from core import msgraphapi as graphapi 
from core import mssentinelapi as defender
#from core import multifactor as mfaauth
from core import dataexplorer as adx
from core.mdr import ManagedDetection 
from core.abuseipdbapi import ipquery
from helper import requestheader as helper
from helper import login
from helper.doc import *
from helper.doc import WordDoc
from helper.excel import WorkBook
from helper import login_snyper as snyper

from config.config import *
from loguru import logger
from docx import Document
import argparse
from datetime import datetime, timedelta, timezone
from dateutil import parser
import pytz    
import json
import configparser
import pandas as pd
from urllib.parse import urlparse

requests.packages.urllib3.disable_warnings()        #Disable requests warning logs

logger.add(f"./logs/{__name__}.log", mode="w", backtrace=True, diagnose=True, level="INFO", filter="ChromeDriver")
#logger.disable(__name__)
dbgPrint = logger
#dbgPrint.disable(__name__)
dbgPrint.disable("core.mssentinelapi")
dbgPrint.disable("core.msgraphapi")
dbgPrint.disable("helper.login")
dbgPrint.disable("core.dataexplorer")

tz = pytz.timezone('Asia/Hong_Kong')

class NriApp:
    classifications = {
        20 : "True Positive",
        10 : "False Positive",
        30 : "Informational, Expected Activity",
        0   : "Not set"
        }

    determinations = {
        140 : "Line of Business Application", 
        130 : "Confirmed Activity",
        120 : "Not Enough Data to Validata",
        110 : "Not Malicious",
        100 : "Malicious User Activity",
        90  : "Phishing",
        80  : "Compromised Account",
        70  : "Multi Stage Attack",
        60  : "Other",
        50  : "Unwanted Software",
        40  : "Security Testing",
        20  : "Malware",
        0   : "Not set"
        }

    source_list = {
        1       : "Endpoint Detection and Response (EDR)",
        2       : "Antivirus", 
        4       : "SmartScreen", 
        32      : "Custom TI",
        512     : "Microsoft Defender for Office 365 (MDO)",
        1024    : "Automated Investigation",
        4096    : "Custom Detection",
        16384   : "Microsoft Defender for Cloud Apps (MCAS)",
        32768   : "Microsoft 365 Defender",
        65536   : "Identity Protection",
        1048576 : "App Governance Policy"
        }

    def __init__(self, args):
        self.args = args
        self.__organization = []
        tracker = {}
#        self.email = self.args.get("email", "")
#        self.output = self.args.get("output",  os.path.abspath(os.path.dirname(__file__)) + "\\output")
       
        
        self.mdr_config()                   #August 22, 2023 - Added MDR Feature
        self.write_config()
        self.read_config()

        if self.email and not hasattr(self, "msgraph"):
            self.msgraph = graphapi.MSGraphApi(self.email).load_session()
#            self.mfa = mfaauth.MultiFactor(self.email).load_session()
            self.sentinel = defender.MSSentinelApi(self.email).load_session()

#            self.mdr = ManagedDetection()

             
    def mdr_config(self):
        if self.args.get("resetmdr"):
            snyper.Login().delete_credentials()
            sys.exit(0)
        else:
            snyper.Login().setup()
            
    def write_config(self):
        self.config = configparser.RawConfigParser()
        self.config.read(os.path.abspath(os.path.dirname(__file__)) + '\config.ini')
        flag = False
        if self.args.get("email"):
            if self.config.has_section("email"):
                flag = True
#                self.config.set("email", "address", self.args["email"])
                self.config['email'] = {"address" : self.args["email"]}
#                self.email = self.args["email"]
            else:
                flag = True
                self.config.add_section("email")
                self.config['email'] = {"address" : self.args["email"]}
#                self.config.set("email", "address", self.args["email"])
#                self.config["email"]["address"] = self.args["email"]
#                self.email = self.args["email"]

        if self.args.get("output"):
            if self.config.has_section("folder"):
                if not os.path.exists(self.args["output"]):
                    os.makedirs(self.output)     
#                self.config.set("folder", "output", self.args["output"])
                self.config["folder"] = {"output" : self.args["output"]}    
                flag = True
            else:
                self.config.add_section('folder')
#                self.config['folder']['output'] = self.args["output"]
                self.config.set("folder", "output", self.args["output"])
                self.output = self.args["output"]
                if not os.path.exists(self.output):
                    os.makedirs(self.output)
                flag = True

#        else:
#            self.config.add_section('folder')
#            self.config['folder']['output'] = os.path.abspath(os.path.dirname(__file__))
#            self.output = os.path.abspath(os.path.dirname(__file__)) + "\\output"

        if self.args.get("companytags"):
            if self.config.has_section("companyTags"): 
                if os.path.exists(self.args["companytags"]):
  #                  self.config.set("companyTags", "path", self.args["companytags"])
                    self.config["companyTags"]["path"] = self.args["companytags"]
                    flag = True
                else:
                    dbgPrint.error("No such file {sample}", sample=self.args["companytags"])
                    sys.exit()
            else:
                self.config.add_section("companyTags")
#                self.config["companyTags"]["path"] = self.args["companytags"]
                if os.path.exists(self.args["companytags"]):
                    self.config["companyTags"] = {"path": self.args["companytags"]}
#                    self.config.set("companyTags", "path", self.args["companytags"])
#                    self.config["companyTags"]["path"] = self.args["companytags"]
                    flag = True
                
#        else:
#            self.config.add_section('companyTags')
#            self.config["companyTags"]["Path"] = os.path.abspath(os.path.dirname(__file__) + "\\companyTags.csv") 
        if flag:

            with open(os.path.abspath(os.path.dirname(__file__)) + '\config.ini', 'w') as configfile:
#            with open('.\config.ini', 'w+') as configfile:
                self.config.write(configfile)
                dbgPrint.info("Config has been setup.")
                exit(0)
           

    def read_config(self):

        if os.path.exists(os.path.abspath(os.path.dirname(__file__)) + '\\config.ini'):
            self.config.read(os.path.abspath(os.path.dirname(__file__)) + '\\config.ini')
            if self.config.has_section('email'):
                self.email = self.config['email']['address']   
            else:
                dbgPrint.error("Email required. Use -e or --email first to store your email address in config.")
                sys.exit()
            if self.config.has_section('folder'):
                self.output = self.config['folder']['output']
            else:
                dbgPrint.error("No output folder yet. Use -o or --output to store the files.")
                sys.exit()

            if self.config.has_section('companyTags'):
                if os.path.exists(self.config["companyTags"]["path"]):
                    df = pd.read_csv(self.config["companyTags"]["path"])
                    df = df.reset_index() 
                    self.tags = [{"Tag": row["Tag"], "Country": row["Country"] , "Company": row["Company"]}  for index, row in df.iterrows()]
                else:
                    dbgPrint.error("{a} does not exist", a=self.config["companyTags"]["path"])
                    sys.exit(1)
            else:
                dbgPrint.error("No company tags csv yet")
                sys.exit(1)

        else:
            dbgPrint.error("Setup configuration first. Use -h --help for help.")
            sys.exit(1)


    def user_summary(self, query):

        user_object = self.msgraph.search_user(query)[0] if self.msgraph.search_user(query) else {}
        temp = {}
        if user_object:
            if user_object["companyName"] != None:
                temp["companyName"] = user_object["companyName"]
            if user_object["country"] != None:
                temp["country"] = user_object["country"]
            if user_object["displayName"] != None:
                temp["name"] = user_object["displayName"]
            if user_object["userPrincipalName"] != None:
                temp["email"] = user_object["userPrincipalName"]
                mfa = self.msgraph.check_mfa_status(user_object["userPrincipalName"])
                if mfa:
                    temp["MFAStatus"] = mfa[0]["isMfaRegistered"]
                else:
                        temp["MFAStatus"] = "Probably not registered"
            if query != None:
                temp["groups"] = self.msgraph.get_user_groups(query)
        return temp

        def ip_summary(self, param, ip_list):
            ip = param.get("senderIP", "")
            if ip:
                return list(set(ip + ip_list))
            else:
                return ip_list

    def file_summary(self, param, hash_list):
        def walk_dict(object, list_object):
            if isinstance(object, dict):
                for k,v in object.items():
                    if isinstance(v, dict):
                        walk_dict(v, list_object)
                    elif isinstance(v, list):
                        for a in v:
                             walk_dict(a, list_object)
                    elif isinstance(v, str) or isinstance(v, int):
                        if k == "hash" or k == "sha256" or k == "sha1":
                            if v.lower() not in list_object:
                                list_object.append(v)
            elif isinstance(object, list):
                for a in object:
                    walk_dict(a, list_object)
            return list_object
      
        hash_list = walk_dict(param, hash_list)
        return hash_list

    def recipient_summary(self, param, hash_list):
        def walk_dict(object, list_object):
            if isinstance(object, dict):
                for k,v in object.items():
                    if isinstance(v, dict):
                        walk_dict(v, list_object)
                    elif isinstance(v, list):
                        for a in v:
                             walk_dict(a, list_object)
                    elif isinstance(v, str) or isinstance(v, int):
                        if k == "recipient":
                            if v.lower() not in list_object:
                                list_object.append(v)
            elif isinstance(object, list):
                for a in object:
                    walk_dict(a, list_object)
            return list_object
      
        hash_list = walk_dict(param, hash_list)
        return hash_list

    def format_userinfo(self, user_object):
        temp = {}

        if user_object["companyName"] != None:

            temp["companyName"] = user_object["companyName"]
            for i in self.tags:
                if i["Tag"].lower() in user_object["companyName"].lower():
                    if user_object["companyName"].lower() not in [z["company"].lower() for z in self.__organization]:
                        self.__organization.append({"company" :user_object["companyName"], "country": i["Country"]})
#                    self.__company_name = user_object["companyName"]
#                    self.__country = i["Country"]
        if user_object["country"] != None:
            temp["country"] = user_object["country"]
        if user_object["usageLocation"] != None:
            temp["usageLocation"] = user_object["usageLocation"]
        if user_object["displayName"] != None:
            temp["name"] = user_object["displayName"] + " (" + user_object["userPrincipalName"] + ")"
        if user_object["userPrincipalName"] != None:
            temp["email"] = user_object["userPrincipalName"]
            mfa = self.msgraph.check_mfa_status(user_object["userPrincipalName"])
            if mfa:
                temp["MFAStatus"] = mfa[0]["isMfaRegistered"]
            else:
                temp["MFAStatus"] = "Probably not registered"
        return temp

    def device_country(self, device_name):
        pass

    def summary(self, param):
        data = {}
        incidentId = param["IncidentId"]
        title = param["Title"]
        tmp_severity = {v:k for k,v in param["AlertsSeveritiesSummary"].items()}
        severity = tmp_severity[max(tmp_severity)] if len(tmp_severity) > 0 else "None"
        categories = param["Categories"]
        firstActivity = param["FirstEventTime"]
        lastActivity = param["LastEventTime"]
        deviceTags = param["IncidentTags"]["DeviceTags"]
        machines = param["ImpactedEntities"]["Machines"]
        users = param["ImpactedEntities"]["Users"]
        mailboxes = param["ImpactedEntities"]["Mailboxes"]
        dSource = param["DetectionSources"]

        source = []
        for i in dSource:
            try:
                source.append(self.source_list[i])
            except:
                dbgPrint.error("Unknown source : {value}", value=i)
                continue

#        computerName = param["ComputerDnsName"]
        users_list = []
        for i in users:
            temp = {}
            displayName = i["DisplayName"]
            userName = i["UserName"]
            userSid = i["UserSid"]
            query = ""
            if displayName != None:
                query = displayName
            elif userName != None:
                query = userName
            elif userSid != None:
                query = userSid
            user_object = self.msgraph.search_user(query)[0] if self.msgraph.search_user(query) else {}
            if user_object:
                temp = self.format_userinfo(user_object)
                if query != None:
                    temp["groups"] = self.msgraph.get_user_groups(query)
                users_list.append(temp)

        device_list = []
        for i in machines:
            temp = {}
            computerName = i["ComputerDnsName"]
            exposureScore = i["ExposureScore"]
            senseMachineId = i["SenseMachineId"]
#            device_object = self.msgraph.search_device_by_name_beta(computerName.split(".")[0])
            device_object = self.msgraph.search_device_by_name(computerName.split(".")[0])[0] if self.msgraph.search_device_by_name(computerName.split(".")[0]) else {}
            if device_object:
                if device_object["deviceName"] != 'none':
                    temp["deviceName"] = device_object["deviceName"]
                    for i in self.tags:
                        if i["Tag"].lower() == device_object["deviceName"].split("-")[0].lower().rstrip('0123456789'):
                            if i["Tag"].upper() not in [z["company"].upper() for z in self.__organization]:
                                self.__organization.append({"company" :i["Tag"].upper(), "country": i["Country"]})
#                            self.__company_name = i["Tag"]
#                            self.__country = i["Country"]
                if device_object["complianceState"] != 'none':
                    temp["complianceState"] = device_object["complianceState"]
                if device_object["azureADDeviceId"] != 'none':
                    temp["azureADDeviceId"] = device_object["azureADDeviceId"]
                if device_object["osVersion"] != 'none':
                    temp["osVersion"] = device_object["osVersion"]
                if device_object["userPrincipalName"] != 'none' and device_object["userPrincipalName"] != '':
                    user_object = self.msgraph.search_user(device_object["userPrincipalName"])[0] if self.msgraph.search_user(device_object["userPrincipalName"]) else {}
                    user = {}
#                    temp["userPrincipalName"] = user
                    if user_object:
                        user = self.format_userinfo(user_object)
    #                temp["userPrincipalName"] = device_object["userPrincipalName"]
                    user["groups"] = self.msgraph.get_user_groups(device_object["userPrincipalName"])
                    temp["Owner"] = user
                device_list.append(temp)
            else:
                device_info = self.sentinel.get_device_info(senseMachineId)
                if device_info:
                    device_object = self.msgraph.search_device_by_azure_id(device_info["AadDeviceId"])[0] if device_info.get("AadDeviceId") and self.msgraph.search_device_by_azure_id(device_info["AadDeviceId"] ) else {}
                    if device_object:
                        if device_object["deviceName"] != 'none':
                            temp["deviceName"] = device_object["deviceName"]
                            for i in self.tags:
                                if i["Tag"].lower() == device_object["deviceName"].split("-")[0].lower().rstrip('0123456789'):
                                    if i["Tag"].upper() not in [z["company"].upper() for z in self.__organization]:
                                        self.__organization.append({"company" :i["Tag"], "country": i["Country"]})
                        if device_object["complianceState"] != 'none':
                            temp["complianceState"] = device_object["complianceState"]
                        if device_object["azureADDeviceId"] != 'none':
                            temp["azureADDeviceId"] = device_object["azureADDeviceId"]
                        if device_object["osVersion"] != 'none':
                            temp["osVersion"] = device_object["osVersion"]
                        if device_object["userPrincipalName"] != 'none':
                            user_object = self.msgraph.search_user(device_object["userPrincipalName"])[0] if self.msgraph.search_user(device_object["userPrincipalName"]) else {}
                            user = {}
        #                    temp["userPrincipalName"] = user
                            if user_object:
                                user = self.format_userinfo(user_object)
            #                temp["userPrincipalName"] = device_object["userPrincipalName"]
                            user["groups"] = self.msgraph.get_user_groups(device_object["userPrincipalName"])
                            temp["Owner"] = user
                        device_list.append(temp)   
                    else:
                        temp["deviceName"] = device_info["ComputerDnsName"]
                        temp["OSPlatform"] = device_info["OsPlatform"]
                        temp["IsAadJoined"] = device_info["IsAadJoined"]
                        device_list.append(temp)
                pass

        mailbox_list = []
        for i in mailboxes:
            temp = {}
            userPrincipalName = i["Upn"]
            user_object = self.msgraph.search_user(userPrincipalName)[0] if self.msgraph.search_user(userPrincipalName) else {}
            if user_object:
                temp = self.format_userinfo(user_object)
                temp["groups"] = self.msgraph.get_user_groups(userPrincipalName)
                mailbox_list.append(temp)
        
        impactedAssets = {"users"       : users_list, 
                         "machines"     : device_list,
                         "mailboxes"    : mailbox_list
                         }


        if deviceTags:
            for i in self.tags:
                if i["Tag"].lower() == deviceTags[0].split("_")[0].lower():
                    if i["Tag"].upper() not in [z["company"].upper() for z in self.__organization]:
                        self.__organization.append({"company" :i["Tag"].upper(), "country": i["Country"]})
                    break

        data["incidentID"]      = incidentId
        data["incidentName"]   = title
        data["severity"]        = severity
        data["CSIRTSeverity"]  = " "
        data["verdict"] = self.classifications[param["Classification"]]
        data["categories"]      = categories
        if self.__organization:
            temp = []
            for z in self.__organization:
                temp.append(z["company"] + " - " +  z["country"])
            data["companyName /Country"] = temp
        else:
            data["companyName /Country"] = "N/A"
#        data["determination"] = self.determinations[param["Determination"]]
        if source:
            data["detectionSource"] = source
        utc_time = parser.parse(firstActivity)
        utc_time = utc_time.replace(tzinfo=pytz.UTC) #replace method      
        ph_time=utc_time.astimezone(tz)        #astimezone method
        data["firstActivity"]   = ph_time.strftime('%Y-%m-%d %H:%M:%S GMT+8')
        utc_time = parser.parse(lastActivity)
        utc_time =utc_time.replace(tzinfo=pytz.UTC) #replace method      
        ph_time=utc_time.astimezone(tz)        #astimezone method
        data["lastActivity"]    = ph_time.strftime('%Y-%m-%d %H:%M:%S GMT+8')
        data["deviceTags"]      = deviceTags
        data["impactedAssets"]  = impactedAssets
        self.__organization = []        
        return data
    
    def ip_summary(self, param, ip_list):
        ip = param.get("senderIP")
        if ip:
            return list(set(ip + ip_list))
        else:
            return ip_list

    def get_audit_logs(self, incident_id):
        start = time.time()
        audit = self.sentinel.get_audit_history_raw(incident_id)
        end = time.time()
        dbgPrint.success("Job done (Elapsed time {value}).", value=end-start)        

        dataexplorer = adx.DataIngestion()
        dataexplorer.ingest_raw_audit_history(data=audit)


    def get_timely_report(self, lookBackInDays, duration=5):
        
        t_end = time.time() + 60 * duration
        dataexplorer = adx.DataIngestion()

        utc_timezone = pytz.UTC

        while time.time() < t_end:

            dbgPrint.debug("Fetching incidents data")
            incidents = self.sentinel.get_incidents(alertStatus=['New','InProgress', 'Resolved'] , 
                                                    severity=[256,128,64,32], 
                                                    pageIndex=1, 
                                                    lookBackInDays=lookBackInDays, 
                                                    pageSize=50, 
                                                    ) #sourceFilter=[16384, 1048576], titleFilter=["eDiscovery"]      #16384 == MCAS incidents 512 = eDiscovery
            dbgPrint.debug("Ingesting data")
            audit = dataexplorer.ingest_events(data=incidents)
            dbgPrint.debug("Fetching audit logs of incidents")

#            audit_logs = self.sentinel.get_all_audit_logs(audit)
#            dbgPrint.debug("Ingesting data")
#            dataexplorer.ingest_audit_logs(data=audit_logs)
#            del audit_logs[:]
            del incidents[:]

            raw_audit_logs = self.sentinel.get_all_audit_logs_raw(new_list=audit, lookBackInDays=lookBackInDays)
            dataexplorer.ingest_raw_audit_history(data=raw_audit_logs)

            if dataexplorer.get_latest_compliance_ingestion_time():
                policy = self.msgraph.get_compliance_policy()
                dataexplorer.ingest_compliance_policy(data=policy) 

            if not dataexplorer.check_device_inventory_tag():
                device_overview = self.sentinel.get_all_device_inventory()
                dataexplorer.ingest_device_inventory(device_overview)

            av_engines = dataexplorer.get_latest_av_engines()
            recent_date = dataexplorer.check_latest_av_engines()
            if not recent_date:
                dataexplorer.ingest_latest_av_engine(data=av_engines)

            #signin_logs = []
            #dataexplorer.ingest_signin_logs(data=signin_logs)
            #earliest_stamp, oldest_stamp = dataexplorer.check_signin_last_timestamp()
            #t_lower = (datetime.utcnow()).replace(minute=0, second=0, microsecond=0, tzinfo=timezone.utc)
            #back_in_time = (t_lower - timedelta(days=lookBackInDays)).replace(minute=0, second=0, microsecond=0)
            #if earliest_stamp:
            #    if (oldest_stamp < back_in_time and back_in_time < earliest_stamp) or (earliest_stamp < back_in_time):
            #        lower_limit = earliest_stamp
            #        upper_limit = None
            #    elif oldest_stamp > back_in_time:                   
            #        lower_limit = back_in_time
            #        upper_limit = oldest_stamp
            #else:
            #    if t_lower.day > 30:
            #        lower_limit = (t_lower - timedelta(days=30)).replace(minute=0, second=0, microsecond=0)
            #    else:
            #        lower_limit = back_in_time
            #last_created_time = self.msgraph.get_all_signins(login_logs=signin_logs, min=lower_limit, max=upper_limit)
            #dataexplorer.ingest_signin_logs(data=signin_logs)
            #del signin_logs[:]

            elif utc_timezone.localize(parser.parse(av_engines["ReleaseDateUtc"])) > parser.parse(recent_date):
                dataexplorer.ingest_latest_av_engine(data=av_engines)


            if not dataexplorer.check_ingest_mfastats_tag():
                mfa_list = self.msgraph.get_users_mfa()
                dataexplorer.ingest_mfastats(data=mfa_list)

            if not dataexplorer.check_ingest_managed_devices_tag():
                dbgPrint.debug("Fetching all Intune devices")
                devices = self.msgraph.get_all_devices_beta()
                dbgPrint.debug("Ingesting data")
                dataexplorer.ingest_managed_devices(data=devices)
                del devices[:]
            
            if not dataexplorer.check_ingest_devices_tags():
                dbgPrint.debug("Fetching managed devices list")
                devices = self.msgraph.get_all_devices()
                dbgPrint.debug("Ingesting data")
                dataexplorer.ingest_devices(data=devices)
                del devices[:]

            if not dataexplorer.check_ingest_users_tag():
                dbgPrint.debug("Fetching all users info")      
                all_users = self.msgraph.get_all_users()
                dbgPrint.debug("Ingesting data")  
                dataexplorer.ingest_users(data=all_users)
                del all_users[:]

            if not dataexplorer.check_defender_agents_tag():
                defender_agents = self.msgraph.get_defender_agents()
                dataexplorer.ingest_defender_agents(path=defender_agents)

            dbgPrint.debug("Sleeping for 3 mins")
            time.sleep(3 * 60)

#        user_list = []
#        mfa_users_list = self.mfa.query_all_users(mfa_list=user_list, page=1)
#        dataexplorer.ingest_mfastats(data=mfa_users_list)   
#        

    def new_template(self, data, doc):
        doc.add_run(str(data["incidentID"]) + " - " + data["incidentName"], True)
        doc.add_run("Severity: " + data["severity"])
        users = data["impactedAssets"]["users"]
        mail = data["impactedAssets"]["mailboxes"]
        machines = data["impactedAssets"]["machines"]
        all_assets = []
        for a in users:
            if a["email"] not in [i["email"] for i in all_assets if i.get("email")]:
                all_assets.append({"email": a["email"], "name": a["name"]})
        for a in mail:
            if a["email"] not in [i["email"] for i in all_assets if i.get("email")]:
                all_assets.append({"email": a["email"],"name": a["name"]})
        for a in machines:
            if a.get("Owner") and a["Owner"].get("email"):
                if a["Owner"]["email"] not in [i["email"] for i in all_assets if i.get("email")]:
                    all_assets.append({"email": a["Owner"]["email"], "name": a["Owner"]["name"], "deviceName" : a["deviceName"]})
                else:
                    for i in all_assets:
                        if i.get("email") and i["email"] == a["Owner"]["email"]:
                            i["deviceName"] = a["deviceName"]
            else:
                all_assets.append({"deviceName": a["deviceName"]})
                    
        doc.add_run("Impacted Assets: ")
        for i in all_assets:
            _ = []
            if i.get("name"):
                _.append(i['name'])
            if i.get("deviceName"):
                _.append(i['deviceName'])
            doc.insert_paragraph(", ".join(_), style="List Bullet 2")
        list_country = ", ".join(data["companyName /Country"])
        doc.add_run("CompanyName / Country: " + list_country)
        doc.add_run("Detection Sources: " + ", ".join(data["detectionSource"]))
        doc.add_run("First Activity: " + data['firstActivity'])
        doc.add_run("Last Activity: " + data["lastActivity"])
        doc.add_run("Verdict:")
        doc.add_run("Incident Summary: ", True)
        doc.add_run("")
        doc.add_run("IOCs:", True)
        doc.add_run("")
        doc.add_run("Alert logs: ", True)
        logs = ["https://security.microsoft.com/alerts/"+i["AlertId"] for i in self.sentinel.get_associated_alerts(str(data["incidentID"]), filter=["eDiscovery"])]
        for i in logs:
            doc.insert_paragraph(i, style="List Bullet")
        doc.add_run("References:", True)
        doc.add_run("")
        pass
    
    def get_full_report_backend(self, incident, lookBackInDays=30, open_doc=False):
        doc = WordDoc()
        ip_list = []
        hash_list = []
        recipient = []
        q = queue.Queue()
        out = self.summary(incident)
        self.new_template(out, doc)

        doc.title(out["incidentName"])
        doc.author(self.email)
        doc.traverse(out, bold=True)

        self.sentinel.get_associated_evidences(incident["IncidentId"], queue=q, lookBackInDays=lookBackInDays)   
        all_data = self.sentinel.accumulate(q)
        for a in all_data:
            doc.insertHR(doc.insert_paragraph())
            ip_list = self.ip_summary(a, ip_list)
            hash_list = self.file_summary(a, hash_list)
            recipient = self.recipient_summary(a, recipient)

            doc.traverse(a)

        doc.insertHR(doc.insert_paragraph())
        doc.add_run("Additional Details")
        doc.traverse([{count: ipquery.check_endpoint(i)} for count, i in enumerate(ip_list) if i != ""])
        doc.insertHR(doc.insert_paragraph())
        doc.traverse([{count: self.sentinel.get_file_info(i)} for count, i in enumerate(hash_list) if i != ""])
        doc.insertHR(doc.insert_paragraph())
        doc.traverse([{count:self.user_summary(i)} for count, i in enumerate(recipient) if i != ""])
        doc.insertHR(doc.insert_paragraph())
        doc.traverse(self.sentinel.get_audit_history(incident["IncidentId"]))
        doc.save(self.output + "\\" + str(incident["IncidentId"]) + ".docx")
        if open_doc:
            os.startfile(self.output + "\\" + str(incident["IncidentId"]) + ".docx")

    def get_full_report(self, id=None, lookBackInDays=30, start_index=0, end_index=0):

        all_data = {}        
        incident_list = self.sentinel.get_incidents(alertStatus=['New','InProgress', 'Resolved'] , 
                                                severity=[256,128,64,32], 
                                                pageIndex=1, 
                                                lookBackInDays=lookBackInDays, 
                                                pageSize=50, 
                                                sourceFilter=[16384, 1048576], titleFilter=["eDiscovery"])       #16384 == MCAS incidents 512 = eDiscovery
        if id:
            incidents = []
            for a in id:
                incident = [i for i in incident_list if i["IncidentId"] == a]
                if incident:
                    incidents.append(incident[0])

        sorted_incidents = sorted(incidents, key=lambda x: x["IncidentId"], reverse=True)
        if sorted_incidents:
            start = sorted_incidents[0]["IncidentId"]
            end = sorted_incidents[len(sorted_incidents)-1]["IncidentId"]
        else:
            dbgPrint.error("Out of range. You may change the lockBackInDays or the incident might be aggragated to other incidents.")
            sys.exit(1)

        if start_index: 
            for count, i in enumerate(sorted_incidents, start=1):
                if start_index >= int(i["IncidentId"]) >= end_index:
                    dbgPrint.info(i["IncidentId"])
                    self.get_full_report_backend(i, lookBackInDays)
                elif int(i["IncidentId"]) < end_index:
                    break
        elif id != None:
            for i in sorted_incidents:
                dbgPrint.info(i["IncidentId"])
                self.get_full_report_backend(i, lookBackInDays, open_doc=True)

        else:
            for count, i in enumerate(sorted_incidents, start=1):
                dbgPrint.info(i["IncidentId"])
                self.get_full_report_backend(i, lookBackInDays)

    def timely_update_mdr(self, lookBackInDays=30):
        dataexplorer = adx.DataIngestion()
        usp = ManagedDetection()

        incident_list = []

        year_list = usp.list_incidents(unit='years', range=lookBackInDays, max=200)
        hour_list = usp.list_incidents(unit='hours', range=24, max=50)

        for i in year_list:
            if i["incidentId"] not in [a["incidentId"] for a in incident_list]:
                incident_list.append(i)

        for i in hour_list:
            if i["incidentId"] not in [a["incidentId"] for a in incident_list]:
                incident_list.append(i)

        day_list = usp.list_incidents(unit='days', range=30, max=50)

        for i in day_list:
            if i["incidentId"] not in [a["incidentId"] for a in incident_list]:
                incident_list.append(i)


        dbgPrint.info("Found " + str(len(incident_list)) + " number of incidents") 
        dataexplorer.ingest_mdr_incidents(data=incident_list)

        policies = []
        all_comments = []
#        incident_list = [{'incidentId':33257627890}, {'incidentId':33257627829},{'incidentId':33257627464}]

        for i in incident_list:

            comments = usp.get_nri_mdr_comments(i['incidentId'])
            if comments:
                all_comments.extend(comments)
#            violator_policy = usp.get_policies_threat(i['violatorId'])
#            if violator_policy:
#                policies.extend(violator_policy)
#            policies.append(usp.get_policies_threat(i['violatorId']))

        dataexplorer.ingest_mdr_audit_logs(data=all_comments)

        pass

    def get_mdr_report(self, id, lookBackInDays=30):

        usp = ManagedDetection()
        self.timely_update_mdr(lookBackInDays=lookBackInDays)
#        usp.get_policies_threat("admin")
#        usp.get_raw_event('33257592814')  #30268684134
#        usp.get_raw_event('30268684134')            #------cannot parse this shit
#        incident_list = usp.list_incidents(unit='years', range=lookBackInDays, max=200)

#        incidents = []
#        for a in id:
#            incident = [i for i in incident_list if i["incidentId"] == a]
#            if incident:
#                incidents.append(incident[0])

#        sorted_incidents = sorted(incidents, key=lambda x: x["incidentId"], reverse=True)
#        sorted_incidents = sorted(incident_list, key=lambda x: x["incidentId"], reverse=True)
        
#        stream = usp.get_incident_activity_stream('26957853650')
#        usp.get_raw_event('26957853650')

#        for i, v in enumerate(sorted_incidents):
#            mdr_comment = usp.get_mdr_comment(v['incidentId'])

#            raw_event = usp.get_raw_event(v['incidentId'])
#            bcd = usp.get_nri_comment(v['incidentId'])
#            abc = usp.get_mdr_comment(v['incidentId'])
#            print(raw_event)
#            xyz = usp.get_policies_threat(v['violatorId'])
#            print(v['incidentId'])                                      #		'incidentId'	'26957853650'	Multiple logs in comments
#            print(v['policyName'])
#            print(i)

#        usp.get_policies_threat("admin")

#        pass

    def dispatcher(self):
        args = self.args
        path = os.path.dirname(os.path.realpath(__file__))

        if args.get("verbose"):
            dbgPrint.enable(__name__)
            dbgPrint.enable("core.mssentinelapi")
            dbgPrint.enable("core.msgraphapi")
            dbgPrint.enable("helper.login")
            dbgPrint.enable("core.dataexplorer")
            dbgPrint.enable("core.mdr")
            dbgPrint.enable("helper.login_snyper")
#        if args.get("user"):
#            self.__init__(args)
#            user_info = self.msgraph.check_mfa_status(args.get("user"))
#            if user_info:
#                printTable(user_info)
#            user_info = self.msgraph.search_user(args.get("user"))
        if args.get("mdr") and args.get("incidentId"):
            dbgPrint.error("Either MDR case or 365 Incident Id")
            sys.exit(1)
        elif args.get("mdr"):
            self.get_mdr_report(args.get("mdr"), lookBackInDays=args.get("daysago"))
        elif args.get("incidentId"):
            self.__init__(args)
            start = time.time()
            self.get_full_report(args.get("incidentId"), lookBackInDays=args.get("daysago") if args.get("daysago") else 5)
            end = time.time()
            dbgPrint.success("Job done (Elapsed time {value}).", value=end-start)
        elif args.get("audit"):
            self.get_audit_logs(args["audit"])
        elif args.get("range"):
            start = time.time()
            dbgPrint.info("Processing Incidents with range")
            start_index = args["range"][0]
            end_index = 0
            if len(args["range"]) > 1 and len(args["range"]) < 3:
                end_index = args["range"][1]
            if (start_index == end_index) or (start_index < end_index):
                dbgPrint.error("The starting index must be higher than ending index")
                sys.exit(1)
            self.get_full_report(lookBackInDays=args.get("daysago") if args.get("daysago") else 30, start_index=start_index, end_index=end_index)
            end = time.time()
            dbgPrint.success("Job done (Elapsed time {value}).", value=end-start)
        elif args.get("timely"):
            if int(args["timely"]):
                self.get_timely_report(lookBackInDays=args.get("daysago") if args.get("daysago") else 5, duration= int(args["timely"]))

            pass

        return

    #------------------------------------------------------

def parse_argument():
        
    parser = argparse.ArgumentParser(
        prog = 'ChromeDriver',
        description  = 'This tool is used to retrieve all the necessary information of the incident',
        epilog = '',
        formatter_class=argparse.RawTextHelpFormatter
        )
#    group = parser.add_mutually_exclusive_group(required=True)

    parser.add_argument('-id', '--incidentId', nargs='+', type=int, help= "Select single or multiple IDs\n"
                                                                          "Example: \n"
                                                                          "-id 13572\n"
                                                                          "-id 13572 13806\n")

    parser.add_argument('-v', '--verbose',  help="Enable logging", action='store_true')
#    parser.add_argument('-r', '--reset', help="Reset the session and restart", action="store_true")
    parser.add_argument('-e', '--email', help="Set email address.")
#    parser.add_argument('-u', '--user', help = "Fetch user info")
    parser.add_argument('-o', '--output', help = "Set output folder.")
#    parser.add_argument('-c','--clear', help = "Clear cache", action="store_true")
    parser.add_argument('-d', '--daysago', help="Look back in days.", default=5, type=int)

    parser.add_argument('-rm', '--resetmdr', help="Reset MDR credential", action='store_true')
 
    parser.add_argument('-c', '--companytags',help="CSV of company tags <Tag> <Company> <Country> to identify which country does a user belong.")

    parser.add_argument('-r', '--range', nargs='+', type=int, help="Specify the range in a reverse order, \n"
                                                                   "wherein the start is the higher limit and end is lower limit.\n"
                                                                   "Lower limit is optional. If not provided, the default value will be zero.\n"
                                                                   "Example: \n"
                                                                   "-r 13700 13000")

    parser.add_argument('-t', '--timely', nargs='?', help="Fetching and ingesting data to Azure Data Explorer and can be use \n" 
                                                          "to run with specific time duration, minimum of 5 mins")

#    parser.add_argument('-f', '--force', nargs='+', help="Force ingest the incident data by using Incident ID and will automatically update the ingested data in Azure Data Explorer")

    parser.add_argument('-a', '--audit', type=int, help="For Audit history of each incidents.\n" 
                        "The input should be incident ID. The only use case of this is if the data \n"
                        " is not be able to ingest automatically, you could use this flag.")

#    parser.add_argument('-m', '--mdr', nargs='+', type=str, help="Ingest MDR tickets")
    parser.add_argument('-m', '--mdr', help="Ingest MDR tickets", action='store_true')

    args = parser.parse_args()

    if len(sys.argv)==1:
        # display help message when no args are passed.
        parser.print_help()
        sys.exit(1)

    return args 

if __name__ == "__main__":

    args = parse_argument()
    myapp = NriApp(vars(args))
    myapp.dispatcher()
#    main(args)